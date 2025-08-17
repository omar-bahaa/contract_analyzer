import os
import logging
from typing import List, Dict, Any, Optional
import fitz  # PyMuPDF
import docx
import pytesseract
from PIL import Image
import io
import re
from camel_tools.tokenizers.word import simple_word_tokenize
from camel_tools.utils.normalize import normalize_alef_maksura_ar, normalize_alef_ar, normalize_teh_marbuta_ar
import pyarabic.araby as araby

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ArabicDocumentProcessor:
    """
    A comprehensive Arabic document processor that handles PDF, DOCX, and image-based text extraction
    with advanced Arabic text normalization and preprocessing.
    """
    
    def __init__(self):
        self.setup_tesseract()
    
    def setup_tesseract(self):
        """Setup Tesseract for Arabic OCR"""
        try:
            # Check if Arabic language pack is available
            langs = pytesseract.get_languages()
            if 'ara' not in langs:
                logger.warning("Arabic language pack not found in Tesseract. OCR quality may be reduced.")
        except Exception as e:
            logger.error(f"Error setting up Tesseract: {e}")
    
    def extract_text_from_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Extract text and metadata from PDF documents"""
        try:
            doc = fitz.open(pdf_path)
            text_content = []
            metadata = {
                'total_pages': len(doc),
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'pages_with_images': []
            }
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Extract text directly
                page_text = page.get_text()
                
                # If text is minimal, try OCR on page images
                if len(page_text.strip()) < 50:
                    try:
                        # Get page as image
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_data))
                        
                        # Perform OCR
                        ocr_text = pytesseract.image_to_string(img, lang='ara+eng')
                        if len(ocr_text.strip()) > len(page_text.strip()):
                            page_text = ocr_text
                            metadata['pages_with_images'].append(page_num + 1)
                    except Exception as e:
                        logger.warning(f"OCR failed for page {page_num + 1}: {e}")
                
                text_content.append({
                    'page': page_num + 1,
                    'text': self.normalize_arabic_text(page_text)
                })
            
            doc.close()
            return {
                'content': text_content,
                'metadata': metadata,
                'full_text': self.combine_pages_text(text_content)
            }
            
        except Exception as e:
            logger.error(f"Error processing PDF {pdf_path}: {e}")
            return {'content': [], 'metadata': {}, 'full_text': ''}
    
    def extract_text_from_docx(self, docx_path: str) -> Dict[str, Any]:
        """Extract text from DOCX documents"""
        try:
            doc = docx.Document(docx_path)
            
            # Extract core properties
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else ''
            }
            
            # Extract text from paragraphs
            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraphs.append({
                        'paragraph': i + 1,
                        'text': self.normalize_arabic_text(para.text)
                    })
            
            # Extract text from tables
            tables_text = []
            for table_idx, table in enumerate(doc.tables):
                table_content = []
                for row_idx, row in enumerate(table.rows):
                    row_content = []
                    for cell in row.cells:
                        cell_text = self.normalize_arabic_text(cell.text)
                        row_content.append(cell_text)
                    table_content.append(row_content)
                
                if any(any(cell.strip() for cell in row) for row in table_content):
                    tables_text.append({
                        'table': table_idx + 1,
                        'content': table_content
                    })
            
            full_text = self.combine_docx_content(paragraphs, tables_text)
            
            return {
                'content': {
                    'paragraphs': paragraphs,
                    'tables': tables_text
                },
                'metadata': metadata,
                'full_text': full_text
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX {docx_path}: {e}")
            return {'content': {'paragraphs': [], 'tables': []}, 'metadata': {}, 'full_text': ''}
    
    def normalize_arabic_text(self, text: str) -> str:
        """Comprehensive Arabic text normalization"""
        if not text:
            return ""
        
        # Basic cleaning
        text = text.strip()
        
        # Remove unwanted characters but keep Arabic punctuation
        text = re.sub(r'[^\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF\s\d\.\,\:\;\!\?\(\)\[\]\-\+\=\/\\\<\>\"\'\`\~\@\#\$\%\^\&\*]', '', text)
        
        # Normalize Arabic characters using CAMeL Tools
        text = normalize_alef_ar(text)
        text = normalize_alef_maksura_ar(text)
        text = normalize_teh_marbuta_ar(text)
        
        # Additional Arabic normalization using pyarabic
        text = araby.strip_harakat(text)  # Remove diacritics
        text = araby.normalize_hamza(text)  # Normalize hamza
        
        # Clean multiple spaces and newlines
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
    
    def extract_clauses_from_contract(self, text: str) -> List[Dict[str, Any]]:
        """
        Extract contract clauses using Arabic-aware patterns
        TODO: edit clause extraction logic to be more comprehensive and smarter
        """
        clauses = []
        
        # Arabic clause indicators
        clause_patterns = [
            r'(?:البند|المادة|الفقرة|القسم)\s*(?:رقم\s*)?(\d+)',
            r'(?:أولاً|ثانياً|ثالثاً|رابعاً|خامساً|سادساً|سابعاً|ثامناً|تاسعاً|عاشراً)',
            r'\d+\.\s*',
            r'\(\d+\)',
            r'[أ-ي]\.\s*'
        ]
        
        # Split text into potential clauses
        sentences = text.split('\n')
        
        clause_num = 1
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 20:  # Filter out very short lines
                # Check if it matches clause patterns
                is_clause = False
                clause_type = "general"
                
                for pattern in clause_patterns:
                    if re.search(pattern, sentence):
                        is_clause = True
                        if 'البند' in sentence:
                            clause_type = "clause"
                        elif 'المادة' in sentence:
                            clause_type = "article"
                        elif 'الفقرة' in sentence:
                            clause_type = "paragraph"
                        break
                
                if is_clause or len(sentence) > 100:  # Include substantial text blocks
                    clauses.append({
                        'id': clause_num,
                        'type': clause_type,
                        'text': sentence,
                        'word_count': len(simple_word_tokenize(sentence))
                    })
                    clause_num += 1
        
        return clauses
    
    def combine_pages_text(self, pages_content: List[Dict]) -> str:
        """Combine text from multiple pages"""
        return '\n\n'.join([page['text'] for page in pages_content if page['text'].strip()])
    
    def combine_docx_content(self, paragraphs: List[Dict], tables: List[Dict]) -> str:
        """Combine paragraphs and tables text"""
        full_text = ""
        
        # Add paragraphs
        for para in paragraphs:
            full_text += para['text'] + '\n\n'
        
        # Add tables
        for table in tables:
            full_text += f"\n--- جدول {table['table']} ---\n"
            for row in table['content']:
                full_text += ' | '.join(row) + '\n'
            full_text += '\n'
        
        return full_text.strip()
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Main method to process any supported document type"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            result = self.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            result = self.extract_text_from_docx(file_path)
        else:
            logger.error(f"Unsupported file type: {file_extension}")
            return {'content': [], 'metadata': {}, 'full_text': '', 'clauses': []}
        
        # Extract clauses if this is a contract
        if result['full_text']:
            result['clauses'] = self.extract_clauses_from_contract(result['full_text'])
        else:
            result['clauses'] = []
        
        return result

# Example usage and testing
if __name__ == "__main__":
    processor = ArabicDocumentProcessor()
    
    # Test with Arabic text normalization
    sample_text = "هذا نص تجريبي للمعالجة العربية مع علامات التشكيل والهمزات المختلفة"
    normalized = processor.normalize_arabic_text(sample_text)
    print(f"Original: {sample_text}")
    print(f"Normalized: {normalized}")
